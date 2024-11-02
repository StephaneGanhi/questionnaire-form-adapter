import pandas as pd
from docx import Document


def adapt_questionnaire_to_teleform(questionnaire_df):
    # Add a column for standardized coding rules
    questionnaire_df['Coding_Rule'] = 'N/A'

    # Define skip logic based on conditions (customizable)
    def add_skip_logic(row):
        if 'skip' in row['Question'].lower():
            return 'Yes'
        else:
            return 'No'
    
    questionnaire_df['Skip_Logic'] = questionnaire_df.apply(add_skip_logic, axis=1)

    # Add a column to highlight terms that need clarification
    def check_for_clarification(row):
        unclear_terms = ['clarify', 'explain', 'specify']
        for term in unclear_terms:
            if term in row['Question'].lower():
                return 'Needs Clarification'
        return 'Clear'
    
    questionnaire_df['Clarification_Needed'] = questionnaire_df.apply(check_for_clarification, axis=1)

    # Ensure format compatibility for teleform (e.g., ASCII text encoding)
    questionnaire_df['Formatted_For_Teleform'] = questionnaire_df['Question'].apply(lambda x: x.encode('ascii', 'ignore').decode('ascii'))

    return questionnaire_df

# Sample data (replace with actual questionnaire data)
data = {
    'Question': [
        'What is your age?',
        'Please specify your country of residence.',
        'Do you want to skip this section?',
        'Explain your current job role in detail.'
    ],
    'Type': ['Numeric', 'Text', 'Yes/No', 'Text']
}

questionnaire_df = pd.DataFrame(data)
adapted_questionnaire_df = adapt_questionnaire_to_teleform(questionnaire_df)

# Save to a Word document with input fields
doc = Document()
doc.add_heading('Adapted Questionnaire for Teleform', level=1)

for _, row in adapted_questionnaire_df.iterrows():
    doc.add_paragraph(f"{row['Question']}")
    if row['Type'].lower() == 'numeric':
        doc.add_paragraph('Answer: ____________________', style='Normal')
    elif row['Type'].lower() == 'text':
        doc.add_paragraph('Answer: ____________________________________________________', style='Normal')
    elif row['Type'].lower() == 'yes/no':
        doc.add_paragraph('Answer: [Yes]   [No]', style='Normal')
    doc.add_paragraph('')  # Add space between questions

doc.save('adapted_teleform_questionnaire.docx')

print("Form-style questionnaire created successfully and saved as 'adapted_teleform_questionnaire.docx'.")
